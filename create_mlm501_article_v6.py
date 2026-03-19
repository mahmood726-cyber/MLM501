import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = docx.Document()

def set_cell_background(cell, fill_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_professional_table(data, header_rows=1, first_col_bold=False):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            paragraphs = cell.paragraphs
            for p in paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(8.5)
                run.font.name = 'Arial'
                if i < header_rows:
                    run.bold = True
                    set_cell_background(cell, "F2F2F2") # Subtle gray for headers
                elif j == 0 and first_col_bold:
                    run.bold = True
    return table

# Title
title = doc.add_heading('MLM501: a multilevel evidence synthesis suite for the Cochrane Fragility Atlas and Clinical Pathway Simulation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Authors
doc.add_heading('Authors', 1)
doc.add_paragraph('Mahmood Ahmad [1], Niraj Kumar [2], Laiba Khan [2], Bilaal Dar [3], Andrew Woo [4]')
doc.add_paragraph('Corresponding author: Mahmood Ahmad (mahmood.ahmad2@nhs.net)')

# Affiliations
doc.add_heading('Affiliations', 1)
doc.add_paragraph('1. National Health Service (NHS), UK\n2. Royal Free Hospital, London, UK\n3. King\'s College London Medical School, London, UK\n4. St George\'s Hospital, London, UK')

# Abstract
doc.add_heading('Abstract', 1)
doc.add_paragraph('Background: Meta-analyses require searching, screening, data entry and analysis in a laborious process. Due to the intensity of this process "living" meta-analysis is difficult to sustain in practice. Previous attempts have either used fully automated or manual methods. MLM501 provides a semi-automated and fully integrated living meta-analysis workflow for evidence synthesis across 501 Cochrane reviews.')
doc.add_paragraph('Methods: It exists as an R package implementing a multilevel effects table (2,373 analyses). It features the Meta-Analysis Fragility Index (MAFI 2.0) engine for robustness auditing, a Bayesian Fragility & Credibility Engine for uncertainty assessment, and an Evidence Pathway Engine for simulating clinical workflows. The package integrates non-linear evidence decay models and health-economic impact metrics.')
doc.add_paragraph('Results: The software successfully executed a global audit of 2,373 clinical meta-analyses, identifying 288 "High Fragility" analyses. A Cardiovascular case study demonstrated a Net Clinical Benefit (NCB) of 163 lives improved per 10,000 patients, illustrating the software\'s ability to bridge diagnostic accuracy and treatment effects.')
doc.add_paragraph('Conclusions: MLM501 is suited for rapid evidence auditing and represents a new synthesis of human and automated methods with rigorous check-gates.')

# Table 1: Visual Abstract
doc.add_heading('Table 1. Visual Abstract: The Cochrane Fragility Atlas Pipeline', 1)
va_data = [
    ['Pipeline Stage', 'Primary Objective', 'Technical Component', 'Strategic Benefit'],
    ['1. Data Ingestion', 'Unified access to 2,373 Cochrane analyses.', 'Multilevel Database Importer', 'Full provenance (DOI/URL) for 501+ reviews.'],
    ['2. Robustness Audit', 'Quantify conclusion stability via stress-tests.', 'MAFI 2.0 Engine', 'Identifies "High Fragility" evidence frontier.'],
    ['3. Pathway Simulation', 'Bridge diagnostic data to treatment harms.', 'DTA/TE Bridge', 'Calculates real-world Net Clinical Benefit.'],
    ['4. Economic Valuation', 'Translate clinical impact into monetary value.', 'QALY Impact Calculator', 'Economic justification for NHS research priority.'],
    ['5. Visual Synthesis', 'Interactive evidence exploration.', 'Glassmorphism Dashboard', 'Transparent analyst and stakeholder interaction.']
]
create_professional_table(va_data)
doc.add_paragraph('The integration of these five stages ensures that evidence synthesis is not merely a summary of existing data but an active audit of its future utility. By bridging the gap between raw data ingestion and economic valuation, MLM501 provides a high-throughput framework for evidence-based decision making.')

# Table 2: Abbreviations & Symbols
doc.add_heading('Table 2. Glossary of Statistical Abbreviations and Symbols', 1)
abbr_data = [
    ['Symbol', 'Definition', 'Functional Role in MLM501'],
    ['MAFI', 'Meta-Analysis Fragility Index', 'Composite score (0–1) measuring conclusion vulnerability.'],
    ['NCB', 'Net Clinical Benefit', 'Final lives-saved metric per 10,000 patients after harms.'],
    ['HKSJ', 'Hartung-Knapp-Sidik-Jonkman', 'Variance adjustment for random-effects uncertainty.'],
    ['QALY', 'Quality-Adjusted Life Year', 'Health-economic unit for Net Economic Value (NEV).'],
    ['logOR', 'Log Odds Ratio', 'Primary effect measure for dichotomous outcomes.'],
    ['SMD', 'Std. Mean Difference', 'Primary effect measure for continuous outcomes.'],
    ['I²', 'Heterogeneity Statistic', 'Used as a penalty in the MAFI robustness calculation.']
]
create_professional_table(abbr_data)
doc.add_paragraph('These symbols form the mathematical foundation of the MLM501 suite. The primary metric, MAFI, is designed to be highly sensitive to the types of data inconsistencies often found in large-scale clinical repositories, allowing for a standardized assessment across diverse medical specialties.')

# Table 3: MAFI 2.0 Weighting Schema
doc.add_heading('Table 3. MAFI 2.0 Composite Score: Weighted Stress-Test Logic', 1)
mafi_data = [
    ['Component', 'Weight', 'Operational Definition', 'Vulnerability Target'],
    ['Direction (DFI)', '30%', 'Does removing one study flip the effect direction?', 'Inference Reversal'],
    ['Significance (SFI)', '25%', 'Does it change the p < 0.05 significance threshold?', 'Statistical Precision'],
    ['Clinical (CFI)', '20%', 'Does the effect cross the defined MCID (e.g. 0.2)?', 'Clinical Relevance'],
    ['Effect Stability (ESI)', '15%', 'Max proportional change in point estimate.', 'Magnitude Robustness'],
    ['CI Stability (CISI)', '10%', 'Consistency of null-exclusion by confidence bounds.', 'Uncertainty Interval Stability'],
    ['Penalties', 'Variable', 'I² > 50% and k < 10 studies.', 'Heterogeneity and Small-k Risks']
]
create_professional_table(mafi_data)
doc.add_paragraph('The weighted components of the MAFI 2.0 index represent a significant evolution in fragility metrics. By accounting for clinical thresholds (MCID) and directionality flips, the software identifies conclusions that are not just statistically precarious but clinically vulnerable.')

# Table 4: Evidence Decay Parameters
doc.add_heading('Table 4. Sigmoidal Evidence Decay: Modeling Evidence "Rot"', 1)
decay_data = [
    ['Parameter', 'Value', 'Mathematical Context', 'Clinical Rationale'],
    ['Midpoint (x0)', '12 Years', 'Sigmoid center point.', 'Evidence utility declines sharply after 12y.'],
    ['Slope (k)', '0.3', 'Rate of utility decay.', 'Assumes rapid technological/clinical drift.'],
    ['"Fresh" Phase', '0–8 Years', 'Utility > 0.8', 'Reliable for current standard-of-care.'],
    ['"Stale" Phase', '> 15 Years', 'Utility < 0.2', 'High risk of outdated methodology/biases.'],
    ['Nesting Inflation', 'log1p(k)/10', 'Correction for multilevel nesting.', 'Shared bias sources in high-volume reviews.']
]
create_professional_table(decay_data)
doc.add_paragraph('Temporal relevance is a critical but often overlooked dimension of evidence synthesis. The implementation of a sigmoidal decay function acknowledges that the clinical half-life of evidence is influenced by the rapid pace of medical innovation, necessitating more frequent updates for older landmark trials.')

# Table 5: Pathway Simulation Parameters
doc.add_heading('Table 5. Input and Output Matrix for Evidence Pathway Simulation', 1)
path_in_data = [
    ['Input Category', 'Key Variables', 'Output Metric', 'Clinical Translation'],
    ['Diagnostic (DTA)', 'Prevalence, Sensitivity, Specificity', 'True Positives (TP) / False Pos. (FP)', 'Diagnostic Accuracy Floor'],
    ['Treatment (TE)', 'Relative Risk (RR), Baseline Risk', 'Events Prevented', 'Biological Efficacy'],
    ['Safety (Harms)', 'Harm Rate (Side effects)', 'Total Treatment Harms', 'Iatrogenic Impact'],
    ['Integrated', 'NCB Formula', 'Net Clinical Benefit', 'Lives Improved per 10k Patients']
]
create_professional_table(path_in_data)
doc.add_paragraph('The Evidence Pathway Engine represents the practical application of these robustness audits. By simulating how diagnostic errors propagate through treatment effects, researchers can visualize the Evidence Frontier—the point at which statistical uncertainty begins to compromise patient safety and net clinical benefit.')

# Table 6: Health Economic Assumptions
doc.add_heading('Table 6. Health Economic Valuation Framework (NHS Benchmarks)', 1)
econ_data = [
    ['Metric', 'Benchmark Value', 'Source/Assumption', 'Use in MLM501'],
    ['Cost per QALY', '£30,000', 'NICE/NHS Threshold', 'Monetary valuation of clinical benefit.'],
    ['QALY Gain/Event', '5.0', 'Stroke/Major MACE prevention', 'Value of one prevented event.'],
    ['QALY Loss/Harm', '0.5', 'Minor clinical adverse event', 'Cost of treatment-related harm.'],
    ['Net Economic Value', 'Formula-driven', 'Benefit Value - Cost of Harm', 'Research priority justification.']
]
create_professional_table(econ_data)
doc.add_paragraph('Economic valuation provides the necessary context for NHS commissioning and research investment. By translating lives saved into Net Economic Value, MLM501 offers a transparent mechanism for justifying the cost of new clinical trials in areas of high evidence fragility.')

# Table 7: Global Robustness Profile (Audit Results)
doc.add_heading('Table 7. Global Evidence Robustness Profile (Audit of 2,373 Meta-analyses)', 1)
audit_data = [
    ['MAFI Classification', 'Count', 'Percentage', 'Clinical Recommendation'],
    ['Robust (MAFI < 0.15)', '400', '16.9%', 'Gold Standard; use for guideline anchors.'],
    ['Low Fragility (0.15–0.30)', '1,219', '51.4%', 'Reliable; minor monitoring recommended.'],
    ['Moderate Fragility (0.30–0.50)', '466', '19.6%', 'Interpret with caution; sensitive to exclusions.'],
    ['High Fragility (MAFI > 0.50)', '288', '12.1%', 'Evidence Frontier; urgent re-verification required.']
]
create_professional_table(audit_data)
doc.add_paragraph('The global audit of 2,373 analyses reveals a significant vulnerability gap in modern medicine. While over half of the audited analyses are classified as robust or low fragility, the identifying of nearly 300 high-fragility cohorts highlights urgent areas for research verification.')

# Table 8: Specialty-Specific Vulnerability
doc.add_heading('Table 8. Specialty-Specific Robustness Distribution (Heuristic Mapping)', 1)
spec_data = [
    ['Medical Specialty', 'Dominant Measure', 'Median MAFI', 'Vulnerability Trend'],
    ['Cardiovascular', 'logOR (Dichotomous)', '0.24', 'Highly robust landmark trials.'],
    ['Neurology/Dementia', 'SMD (Continuous)', '0.48', 'High volatility in small-k cohorts.'],
    ['Infectious Disease', 'logOR / RR', '0.31', 'Moderate stability; sensitive to recent RCTs.'],
    ['Respiratory', 'MD (Continuous)', '0.39', 'Higher fragility due to heterogeneity (I²).'],
    ['General Medicine', 'Varied', '0.35', 'Mixed stability profile across 501 reviews.']
]
create_professional_table(spec_data)
doc.add_paragraph('Specialty-specific distributions highlight the varying evidence landscapes within clinical medicine. For instance, the relative stability of Cardiovascular trials contrasts with the higher volatility observed in Neurology, likely reflecting differences in sample sizes and the maturity of standardized outcome measures.')

# Table 9: Internal Stress-Test Validation (Unit Test Scenarios)
doc.add_heading('Table 9. Software Validation: Internal Stress-Test Performance', 1)
val_data = [
    ['Test Scenario', 'Target Logic', 'Expected Pass Criteria', 'Observed Status'],
    ['Single-Study Flip', 'MAFI Directionality', 'Detection of inference reversal.', 'PASS'],
    ['P-value Threshold', 'Significance (SFI)', 'Correct handling of p=0.05 boundary.', 'PASS'],
    ['Missing SE Check', 'Cohort Filtering', 'Finite/positive seTE enforcement.', 'PASS'],
    ['Bayesian Parity', 'Posterior Stability', '95% CrI containment of freq. estimate.', 'PASS (diff < 0.01)']
]
create_professional_table(val_data)
doc.add_paragraph('Internal validation ensures that the statistical logic remains consistent across edge cases. The software\'s ability to handle missing standard errors and detect inference reversals has been stress-tested across multiple clinical scenarios, achieving high-precision parity with established R libraries.')

# Table 10: Submission Checklist & Data Availability
doc.add_heading('Table 10. Software Availability and Repository Checklist', 1)
check_data = [
    ['Item', 'Current Status', 'Location/Format', 'Access Type'],
    ['Software Source', 'Complete (v0.1.0)', 'R Package (sub-directory)', 'Open Source (MIT)'],
    ['Multilevel Database', '2,373 analyses', 'data-raw/ (read-only)', 'Restricted (provenance preserved)'],
    ['Guided Vignette', 'Available', 'vignettes/CFA_Walkthrough.Rmd', 'Documentation'],
    ['Audit Ledger', 'Available', 'inst/extdata/cfa_comprehensive_audit.csv', 'CSV Artifact'],
    ['Visual Dashboard', 'Functional', 'inst/extdata/cfa_dashboard.html', 'Glassmorphism UI']
]
create_professional_table(check_data)
doc.add_paragraph('The software availability checklist confirms our commitment to open science and reproducibility. By providing the source code, guide vignettes, and the full audit ledger, we enable independent verification of our findings and support the wider research community in adopting dynamic evidence auditing.')

# Discussion
doc.add_heading('Discussion', 1)
doc.add_paragraph('The MLM501 suite represents a paradigm shift from Snapshot meta-analysis to Dynamic Auditing. Traditional systematic reviews are often outdated before publication; by automating the audit of 2,373 analyses, we enable real-time monitoring of global clinical knowledge. The integration of the Evidence Pathway Engine further elevates this work by identifying the Cost of Fragility—the economic and clinical price of making decisions based on unstable data. This framework allows the NHS to move beyond simple effect sizes toward a more nuanced Return on Evidence model for research prioritization.')

# Limitations
doc.add_heading('Limitations', 1)
doc.add_paragraph('Despite its robustness, several limitations remain. The specialty mapping relies on heuristic Cochrane ID sequences, which may misclassify multi-disciplinary reviews. Additionally, the computational cost of Stan-based Bayesian simulations limits their application to only the most fragile cohorts. Future versions will aim to integrate real-time API-driven specialty tagging and more efficient approximation methods for posterior estimation.')

# Conclusion
doc.add_heading('Conclusion', 1)
doc.add_paragraph('MLM501 provides a groundbreaking framework for the Cochrane Fragility Atlas, enabling the first-of-its-kind global audit of clinical evidence robustness. By identifying the Vulnerability Frontier and simulating real-world pathway impacts, it offers an essential tool for healthcare researchers and policy-makers to optimize evidence-based care and research investment.')

# Figure Legends
doc.add_heading('Figures and Visual Walkthrough', 1)
doc.add_paragraph('Note: Figures are provided as separate high-resolution PNG files.')
doc.add_paragraph('Figure 1. Entropy Landscape: Mapping MAFI vs. Heterogeneity (I²). (cfa_entropy_landscape.png)')
doc.add_paragraph('Figure 2. Multi-Measure Landscape: Global distribution across logOR, SMD, and GIV. (cfa_multi_measure_landscape.png)')
doc.add_paragraph('Figure 3. Temporal Trends in Evidence Fragility: The Sigmoidal Evidence Rot. (cfa_temporal_trends.png)')
doc.add_paragraph('Figure 4. Cardiovascular Pathway Frontier: Sensitivity threshold mapping. (pathway_frontier_cvd.png)')
doc.add_paragraph('Figure 5. Global Metric Distribution: Statistical coverage of 2,373 analyses. (cfa_measure_distribution.png)')

# References
doc.add_heading('References', 1)
doc.add_paragraph('1. DerSimonian R, Laird N. Meta-analysis in clinical trials. Controlled Clinical Trials. 1986.\n2. Ahmad M et al. The Meta-Analysis Fragility Index 2.0. NHS Technical Series. 2026.\n3. IntHout J et al. The Hartung-Knapp-Sidik-Jonkman method... BMC Med Res Methodol. 2014.')

doc.save(r'C:\Users\user\OneDrive - NHS\Documents\501MLM_Submission\F1000_MLM501_Software_Article_v6.docx')
print('Detailed F1000 Article (v6) with expanded text and sections generated successfully.')
