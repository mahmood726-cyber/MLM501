import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = docx.Document()

def set_cell_background(cell, fill_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(data, header_rows=1, first_col_bold=False):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            paragraphs = cell.paragraphs
            for p in paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(9)
                if i < header_rows:
                    run.bold = True
                    set_cell_background(cell, "EDF2F7")
                elif j == 0 and first_col_bold:
                    run.bold = True
    return table

# Title
title = doc.add_heading('MLM501: a multilevel evidence synthesis suite for the Cochrane Fragility Atlas and Clinical Pathway Simulation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Authors
doc.add_heading('Authors', 1)
doc.add_paragraph('Mahmood Ahmad [1]')
doc.add_paragraph('Corresponding author: Mahmood Ahmad (mahmood.ahmad2@nhs.net)')

# Abstract
doc.add_heading('Abstract', 1)
p = doc.add_paragraph()
p.add_run('Background: ').bold = True
p.add_run('Modern evidence synthesis requires more than static meta-analysis; it demands a continuous audit of evidence robustness. The MLM501 R package was developed to manage a multilevel treatment effect database derived from over 500 Cochrane reviews and to provide an automated "Stress-Test" engine for global clinical knowledge.')

p = doc.add_paragraph()
p.add_run('Methods: ').bold = True
p.add_run('MLM501 is an R package implementing a multilevel meta-analysis database of 2,373 unique analyses. It features the Meta-Analysis Fragility Index (MAFI 2.0) engine, a Bayesian Fragility & Credibility Engine, and an Evidence Pathway Engine for simulating clinical workflows with health-economic (QALY) impact metrics.')

p = doc.add_paragraph()
p.add_run('Results: ').bold = True
p.add_run('A global audit identified 288 "High Fragility" analyses vulnerable to single-study changes. A Cardiovascular case study using the Pathway Engine demonstrated a Net Clinical Benefit (NCB) of 163 lives improved per 10,000 patients.')

p = doc.add_paragraph()
p.add_run('Conclusions: ').bold = True
p.add_run('MLM501 provides a groundbreaking framework for auditing evidence integrity and modeling clinical impact, identifying the "Cost of Fragility" for healthcare optimization.')

# Visual Abstract
doc.add_heading('Visual Abstract (Table-Only)', 1)
va_data = [
    ['VISUAL ABSTRACT | Cochrane Fragility Atlas', 'VISUAL ABSTRACT | Cochrane Fragility Atlas', 'VISUAL ABSTRACT | Cochrane Fragility Atlas', 'VISUAL ABSTRACT | Cochrane Fragility Atlas'],
    ['1. Audit Database', '2. Stress-Test Engine', '3. Pathway Simulation', '4. Economic Impact'],
    ['2,373 multilevel Cochrane analyses', 'MAFI 2.0 + Sigmoidal evidence decay', 'Test -> Treat "NCB" calculations', 'QALY-based monetary impact metrics'],
    ['Clarity', 'Balance', 'Verification', 'Benefit'],
    ['Vulnerability map for NHS research priorities', 'Fragility vs. Heterogeneity mapped together', 'Internal Stress-Test & Bayesian parity checks', 'Evidence-based healthcare optimization tool']
]
create_table(va_data)

# Core Architecture Table
doc.add_heading('Internal Architecture', 1)
arch_data = [
    ['Layer', 'Components', 'Role'],
    ['Data Ingestion', 'Multilevel Database Importer', 'Unified access to 2,373 analyses with provenance.'],
    ['Audit Engine', 'MAFI 2.0 + brms/Stan', 'Weighted stress-test + Bayesian credible intervals.'],
    ['Pathway Engine', 'DTA/TE Bridge', 'Simulates NCB and total harms for medical strategies.'],
    ['Health Economics', 'QALY Impact Calculator', 'Translates clinical outcomes into monetary value.'],
    ['Analytics', 'Interactive Glassmorphism Dashboard', 'Global evidence exploration and frontier mapping.']
]
create_table(arch_data)
doc.add_paragraph('Table 4. MLM501 internal architecture.', style='Caption')

# Vignette Reference
doc.add_heading('Software Operation and Vignettes', 1)
doc.add_paragraph('The software includes a comprehensive vignette, "Mapping the Evidence Frontier with MLM501" (vignettes/CFA_Walkthrough.Rmd), which provides a guided workflow from global database auditing to clinical pathway frontier mapping.')

# Global Audit Result Table
doc.add_heading('Global Evidence Robustness Audit Results', 1)
audit_data = [
    ['Classification', 'Count (k=2,373)', 'Clinical Interpretation'],
    ['Robust', '400', 'Conclusions highly stable; gold-standard anchors.'],
    ['Low Fragility', '1,219', 'Generally reliable; low sensitivity to study removal.'],
    ['Moderate Fragility', '466', 'Sensitive to study exclusions; interpret with caution.'],
    ['High Fragility', '288', 'Vulnerable to a single study change; evidence frontier.']
]
create_table(audit_data)
doc.add_paragraph('Table 5. Global robustness profile of Cochrane clinical evidence.', style='Caption')

# Pathway Simulation Table
doc.add_heading('Pathway Simulation Case Study (CVD)', 1)
path_data = [
    ['Metric', 'Simulation Result', 'Economic Context'],
    ['Prevalence', '15%', 'High-risk cardiovascular population.'],
    ['Net Clinical Benefit (NCB)', '163 lives per 10k', 'Improved outcomes after accounting for harms.'],
    ['Total Treatment Harms', '22 events per 10k', 'Adverse events in false/true positives.'],
    ['Net Economic Value', '£4.89M per 10k', 'Based on £30,000 per QALY (NHS benchmark).'],
    ['Evidence Age (Median)', '14 years', 'Stale score adjusted via sigmoidal decay.']
]
create_table(path_data)
doc.add_paragraph('Table 6. Cardiovascular pathway simulation using HS-Troponin -> Antihypertensive therapy.', style='Caption')

# Software availability
doc.add_heading('Software availability', 1)
doc.add_paragraph('Source code: MLM501 sub-directory (GitHub/Zenodo archive)')
doc.add_paragraph('Version: v0.1.0 (The Cochrane Fragility Atlas Suite)')
doc.add_paragraph('License: MIT')

doc.save(r'C:\Users\user\OneDrive - NHS\Documents\501MLM_Submission\F1000_MLM501_Software_Article.docx')
print('F1000 Article generated successfully.')
