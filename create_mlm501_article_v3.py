import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = docx.Document()

def set_cell_background(cell, fill_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(data, header_rows=1, first_col_bold=False):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            paragraphs = cell.paragraphs
            for p in paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(9)
                if i < header_rows:
                    run.bold = True
                    set_cell_background(cell, "EDF2F7")
                elif j == 0 and first_col_bold:
                    run.bold = True
    return table

# Title
title = doc.add_heading('MLM501: a multilevel evidence synthesis suite for the Cochrane Fragility Atlas and Clinical Pathway Simulation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Authors
doc.add_heading('Authors', 1)
doc.add_paragraph('Mahmood Ahmad [1]')
doc.add_paragraph('Corresponding author: Mahmood Ahmad (mahmood.ahmad2@nhs.net)')

# Affiliations
doc.add_heading('Affiliations', 1)
doc.add_paragraph('1. National Health Service (NHS), UK')

# Abstract
doc.add_heading('Abstract', 1)
p = doc.add_paragraph()
p.add_run('Background: ').bold = True
p.add_run('Modern evidence synthesis requires more than static meta-analysis; it demands a continuous audit of evidence robustness and its translation into clinical care pathways. The MLM501 R package was developed to manage a multilevel treatment effect database derived from over 500 Cochrane reviews and to provide an automated "Stress-Test" engine for global clinical knowledge.')

p = doc.add_paragraph()
p.add_run('Methods: ').bold = True
p.add_run('MLM501 is an R package implementing a multilevel meta-analysis database of 2,373 unique analyses. It features the Meta-Analysis Fragility Index (MAFI 2.0) engine, a Bayesian Fragility & Credibility Engine, and an Evidence Pathway Engine for simulating clinical workflows with health-economic (QALY) impact metrics. The package incorporates sigmoidal evidence decay models and specialty-specific robustness mapping.')

p = doc.add_paragraph()
p.add_run('Results: ').bold = True
p.add_run('The software successfully executed a global audit of 2,373 clinical meta-analyses, identifying 288 "High Fragility" analyses vulnerable to single-study changes. A Cardiovascular case study using the Pathway Engine demonstrated a Net Clinical Benefit (NCB) of 163 lives improved per 10,000 patients, illustrating the software\'s ability to bridge diagnostic accuracy and treatment effects.')

p = doc.add_paragraph()
p.add_run('Conclusions: ').bold = True
p.add_run('MLM501 provides a groundbreaking framework for auditing evidence integrity and modeling clinical impact. It is an essential asset for healthcare researchers and policy-makers focused on evidence-based optimization, identifying the "Cost of Fragility" in healthcare.')

# Visual Abstract
doc.add_heading('Visual Abstract (Table-Only)', 1)
va_data = [
    ['VISUAL ABSTRACT | Cochrane Fragility Atlas', 'VISUAL ABSTRACT | Cochrane Fragility Atlas', 'VISUAL ABSTRACT | Cochrane Fragility Atlas', 'VISUAL ABSTRACT | Cochrane Fragility Atlas'],
    ['1. Audit Database', '2. Stress-Test Engine', '3. Pathway Simulation', '4. Economic Impact'],
    ['2,373 multilevel Cochrane analyses', 'MAFI 2.0 + Sigmoidal evidence decay', 'Test -> Treat "NCB" calculations', 'QALY-based monetary impact metrics'],
    ['Clarity', 'Balance', 'Verification', 'Benefit'],
    ['Vulnerability map for NHS research priorities', 'Fragility vs. Heterogeneity mapped together', 'Internal Stress-Test & Bayesian parity checks', 'Evidence-based healthcare optimization tool']
]
create_table(va_data)

# Introduction
doc.add_heading('Introduction', 1)
doc.add_paragraph('The increasing volume of clinical evidence necessitates automated systems to monitor the stability of medical conclusions. Traditional meta-analyses often fail to account for the "fragility" of their findings—the sensitivity of results to the exclusion of a single trial. Furthermore, the translation of these pooled estimates into real-world clinical benefits remains a manual and error-prone process.')
doc.add_paragraph('The MLM501 package addresses these challenges by consolidating 501+ Cochrane Systematic Reviews into a multilevel database and providing an automated audit framework: the Cochrane Fragility Atlas (CFA). This manuscript describes the architecture of MLM501, its implementation of the MAFI stress-test logic, and the Evidence Pathway Engine used to simulate health-economic outcomes.')

# Methods
doc.add_heading('Methods', 1)
doc.add_heading('Implementation & Architecture', 2)
doc.add_paragraph('The software is implemented as an R package leveraging the metafor engine for frequentist meta-analysis and brms (Stan) for Bayesian credible interval estimation. It maintains a read-only multilevel database that preserves the provenance (DOI/URL) of 2,373 analyses across dichotomous, continuous, and generic inverse-variance (GIV) measures.')

arch_data = [
    ['Layer', 'Components', 'Role'],
    ['Data Ingestion', 'Multilevel Database Importer', 'Unified access to 2,373 analyses with provenance.'],
    ['Audit Engine', 'MAFI 2.0 + brms/Stan', 'Weighted stress-test + Bayesian credible intervals.'],
    ['Pathway Engine', 'DTA/TE Bridge', 'Simulates NCB and total harms for medical strategies.'],
    ['Health Economics', 'QALY Impact Calculator', 'Translates clinical outcomes into monetary value.'],
    ['Analytics', 'Interactive Dashboard', 'Global evidence exploration and frontier mapping.']
]
create_table(arch_data)
doc.add_paragraph('Table 4. MLM501 internal architecture.', style='Caption')

doc.add_heading('The Meta-Analysis Fragility Index (MAFI 2.0)', 2)
doc.add_paragraph('The MAFI 2.0 score is a composite metric designed to capture the stability of meta-analytic conclusions across five weighted components:')
doc.add_paragraph('\u2022 Direction (30%): Does removing a study flip the effect direction?\n\u2022 Significance (25%): Does it change the p-value threshold (0.05)?\n\u2022 Clinical (20%): Does it cross the Minimal Clinically Important Difference (MCID)?\n\u2022 Effect Stability (15%): Proportional change in the point estimate.\n\u2022 CI Stability (10%): Stability of confidence interval bounds.')

doc.add_heading('Non-Linear Evidence Decay', 2)
doc.add_paragraph('Addressing the "rotting" of clinical evidence over time, MLM501 implements a sigmoidal decay function: S = 1 / (1 + exp(-0.3 * (Age - 12))). This model assumes evidence stays relatively "fresh" for 5-8 years, then decays rapidly as new technologies and clinical practices emerge, with a midpoint at 12 years.')

doc.add_heading('Evidence Pathway Engine', 2)
doc.add_paragraph('The Pathway Engine bridges Diagnostic Test Accuracy (DTA) and Treatment Effects (TE) using the simulate_pathway() module. It calculates the Net Clinical Benefit (NCB) by modeling True Positives successfully treated against Treatment Harms in false positives and true positives. Economic impact is further quantified using QALY metrics (e.g., £30,000 per QALY).')

# Operation
doc.add_heading('Operation', 1)
doc.add_paragraph('The software includes a comprehensive vignette, "Mapping the Evidence Frontier with MLM501," providing a guided workflow from global database auditing to clinical pathway frontier mapping. Sequential user actions include:')
doc.add_paragraph('1. Loading the comprehensive audit database.\n2. Executing global audits to compute MAFI scores and evidence age.\n3. Applying the Bayesian engine for high-priority stress-testing.\n4. Simulating clinical pathways to model NCB and economic value.')

# Use Cases
doc.add_heading('Use Cases', 1)
doc.add_heading('Case Study 1: Global Evidence Audit (The CFA)', 2)
audit_data = [
    ['Classification', 'Count (k=2,373)', 'Clinical Interpretation'],
    ['Robust', '400', 'Conclusions highly stable; gold-standard anchors.'],
    ['Low Fragility', '1,219', 'Generally reliable; low sensitivity to study removal.'],
    ['Moderate Fragility', '466', 'Sensitive to study exclusions; interpret with caution.'],
    ['High Fragility', '288', 'Vulnerable to a single study change; evidence frontier.']
]
create_table(audit_data)
doc.add_paragraph('Table 5. Global robustness profile of Cochrane clinical evidence.', style='Caption')

doc.add_heading('Case Study 2: Cardiovascular Pathway Simulation', 2)
path_data = [
    ['Metric', 'Simulation Result', 'Economic Context'],
    ['Prevalence', '15%', 'High-risk cardiovascular population.'],
    ['Net Clinical Benefit (NCB)', '163 lives per 10k', 'Improved outcomes after accounting for harms.'],
    ['Total Treatment Harms', '22 events per 10k', 'Adverse events in false/true positives.'],
    ['Net Economic Value', '£4.89M per 10k', 'Based on £30,000 per QALY (NHS benchmark).']
]
create_table(path_data)
doc.add_paragraph('Table 6. Cardiovascular pathway simulation metrics.', style='Caption')

# Figures and Walkthrough
doc.add_heading('Figures and Visual Walkthrough', 1)
doc.add_paragraph('Note: Figures are provided separately as high-resolution PNG files.')

doc.add_paragraph('Figure 1. Entropy Landscape: Mapping MAFI vs. Heterogeneity (I²). This plot illustrates the relationship between evidence fragility and between-study inconsistency across 2,373 analyses.')

doc.add_paragraph('Figure 2. Multi-Measure Landscape. Global map of evidence robustness across dichotomous (logOR), continuous (SMD/MD), and generic inverse-variance (GIV) measures.')

doc.add_paragraph('Figure 3. Temporal Trends in Evidence Fragility. This chart demonstrates the sigmoidal "rotting" of evidence over time, highlighting the vulnerability of conclusions older than 12 years.')

doc.add_paragraph('Figure 4. Cardiovascular Pathway Frontier. Sensitivity threshold mapping for clinical benefit, demonstrating how diagnostic accuracy ripples through to treatment harms.')

doc.add_paragraph('Figure 5. Distribution of Evidence Measures. Summary of the statistical metrics (logOR, SMD, GIV) covered in the global audit.')

# Discussion
doc.add_heading('Discussion', 1)
doc.add_paragraph('MLM501 represents a shift from "Snapshot" meta-analysis to "Dynamic Audit." Its main strength is the integration of evidence robustness with clinical pathway modeling. By identifying the "Cost of Fragility" in economic terms, it provides a powerful argument for research investment. Limitations include the heuristic specialty mapping and the computational intensity of Stan-based Bayesian simulations.')

# Conclusions
doc.add_heading('Conclusions', 1)
doc.add_paragraph('MLM501 provides a groundbreaking framework for the Cochrane Fragility Atlas, enabling real-time auditing of evidence integrity and its translation into clinical value. It is an essential asset for healthcare researchers and policy-makers focused on evidence-based optimization.')

# Software availability
doc.add_heading('Software availability', 1)
doc.add_paragraph('Source code: MLM501 sub-directory (GitHub/Zenodo archive)')
doc.add_paragraph('Version: v0.1.0')
doc.add_paragraph('License: MIT')

# References
doc.add_heading('References', 1)
doc.add_paragraph('1. DerSimonian R, Laird N. Meta-analysis in clinical trials. Controlled Clinical Trials. 1986.\n2. Gelman A, et al. Prior distributions for variance parameters in hierarchical models. Bayesian Analysis. 2006.\n3. IntHout J, et al. The Hartung-Knapp-Sidik-Jonkman method... BMC Med Res Methodol. 2014.\n4. Ahmad M. The Meta-Analysis Fragility Index (MAFI). NHS Technical Report. 2026.\n5. Page MJ, et al. The PRISMA 2020 statement. BMJ. 2021.')

doc.save(r'C:\Users\user\OneDrive - NHS\Documents\501MLM_Submission\F1000_MLM501_Software_Article_v3.docx')
print('Detailed F1000 Article (Text-Only) generated successfully.')
